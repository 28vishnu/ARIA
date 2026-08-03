from pathlib import Path
import logging

from docx import Document as DocxDocument

from ..models import Document, DocumentPage

logger = logging.getLogger("aria")


class DOCXParser:

    async def parse(self, file_path: str) -> Document:

        doc = DocxDocument(file_path)

        document = Document(
            name=Path(file_path).name,
            path=file_path,
            extension=".docx",
            size=Path(file_path).stat().st_size,
        )

        paragraphs = []

        for paragraph in doc.paragraphs:

            text = paragraph.text.strip()

            if text:
                paragraphs.append(text)

        document.pages.append(
            DocumentPage(
                number=1,
                text="\n".join(paragraphs),
            )
        )

        logger.info(
            "[DOCXParser] Parsed %d paragraphs from %s",
            len(paragraphs),
            document.name,
        )

        return document
